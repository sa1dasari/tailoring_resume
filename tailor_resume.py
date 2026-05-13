"""
Tailor a resume to a job description using the Anthropic API.

Setup:
    1. Install dependencies:  python -m pip install -r requirements.txt
    2. Put your resume in inputs/resume.docx (or switch RESUME_PATH to a .docx file).
    3. Put the target JD in inputs/job_description.txt.
    4. Edit rules in rules/resume_rules.txt.
    5. Store API key in a local .env file (not committed):
       ANTHROPIC_API_KEY=sk-ant-...
    6. Run: python tailor_resume.py

Usage notes:
    - Input files are read from inputs/ by default.
    - Output is written to output/tailored_resume_<Company>.docx
"""

import os
import re
import sys
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path

try:
    import anthropic
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ModuleNotFoundError as exc:
    missing_pkg = exc.name or "required package"
    sys.exit(
        f"Missing dependency: {missing_pkg}.\n"
        "Install project dependencies first: python -m pip install -r requirements.txt"
    )

# ─── CONFIG ──────────────────────────────────────────────────────────────────
BASE_DIR       = Path(__file__).resolve().parent
INPUT_DIR      = BASE_DIR / "inputs"
RESUME_PATH    = INPUT_DIR / "Sawan_Dasari_Resume.docx"
JOB_DESC_PATH  = INPUT_DIR / "job_description.txt"
OUTPUT_DIR     = BASE_DIR / "output"              # Generated files are saved here
DOTENV_PATH    = BASE_DIR / ".env"
RULES_PATH     = BASE_DIR / "rules" / "resume_rules.txt"
MODEL          = "claude-sonnet-4-5"   # or "claude-opus-4-5" for the most capable
LINKEDIN_URL   = "https://www.linkedin.com/in/sawan-dasari/"

# Common technologies/tools used for lightweight JD gap detection.
TECH_TERMS = [
    "python", "typescript", "javascript", "java", "go", "c++", "c#", "sql", "nosql",
    "fastapi", "django", "flask", "node.js", "node", "react", "postgres", "mysql", "mongodb",
    "redis", "kafka", "airflow", "docker", "kubernetes", "aws", "gcp", "azure", "terraform",
    "ci/cd", "rest", "graphql", "api", "sdk", "observability", "data pipeline", "integrations",
    "schema", "data modeling", "strong typing", "typing", "libraries", "tooling",
]
# ─────────────────────────────────────────────────────────────────────────────


# ─── INPUT READING ───────────────────────────────────────────────────────────

def extract_docx_text(path: str | Path) -> str:
    """Pull paragraph text from a .docx file."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def load_dotenv_file(path: Path) -> None:
    """Load key=value pairs from a local .env file into environment variables."""
    if not path.exists():
        return

    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and value and key not in os.environ:
            os.environ[key] = value


def load_resume_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return extract_docx_text(path)
    return path.read_text(encoding="utf-8").strip()


def load_inputs() -> tuple[str, str]:
    if not RESUME_PATH.exists():
        sys.exit(f"Resume not found: {RESUME_PATH}")
    if not JOB_DESC_PATH.exists():
        sys.exit(f"Job description not found: {JOB_DESC_PATH}")

    resume_text = load_resume_text(RESUME_PATH)
    if not resume_text:
        sys.exit(f"Resume file is empty: {RESUME_PATH}")

    with open(JOB_DESC_PATH, "r", encoding="utf-8") as f:
        jd_text = f.read().strip()

    if not jd_text:
        sys.exit("Job description file is empty.")

    return resume_text, jd_text


# ─── API CALL ────────────────────────────────────────────────────────────────

def load_system_prompt(path: Path) -> str:
    """Load system prompt/rules from an external text file."""
    if not path.exists():
        sys.exit(f"Rules file not found: {path}")

    prompt = path.read_text(encoding="utf-8").strip()
    if not prompt:
        sys.exit(f"Rules file is empty: {path}")
    return prompt


def _contains_term(text: str, term: str) -> bool:
    txt = text.lower()
    t = term.lower().strip()
    if not t:
        return False
    if re.search(r"[+#./-]", t) or " " in t:
        return t in txt
    return re.search(rf"\b{re.escape(t)}\b", txt) is not None


def extract_missing_jd_terms(resume_text: str, jd_text: str) -> list[str]:
    """Find likely required JD technologies/tools that are absent from the original resume text."""
    in_jd = [term for term in TECH_TERMS if _contains_term(jd_text, term)]
    missing = [term for term in in_jd if not _contains_term(resume_text, term)]
    # Keep deterministic, short guidance to avoid over-constraining the model.
    return sorted(set(missing))[:12]


def call_claude(
    resume_text: str,
    jd_text: str,
    api_key: str,
    system_prompt: str,
    missing_jd_terms: list[str] | None = None,
) -> tuple[str, dict]:
    client = anthropic.Anthropic(api_key=api_key)

    print(f"Calling {MODEL}...")
    skill_guidance = ""
    if missing_jd_terms:
        skill_guidance = (
            "\n\nMISSING_JD_TERMS_POSSIBLY_RELEVANT:\n"
            f"{', '.join(missing_jd_terms)}\n"
            "If any term above can be truthfully inferred from the candidate's existing experience context, "
            "you may weave it into EXPERIENCE bullets. If not, omit it. "
            "Never add unsupported claims, metrics, projects, or responsibilities."
        )

    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=[{
            "role": "user",
            "content": f"RESUME:\n{resume_text}\n\nJOB DESCRIPTION:\n{jd_text}{skill_guidance}"
        }],
    )

    text = response.content[0].text
    usage = {
        "input_tokens": response.usage.input_tokens,
        "output_tokens": response.usage.output_tokens,
    }
    return text, usage


# ─── PARSING ─────────────────────────────────────────────────────────────────

SECTION_HEADERS = {"SUMMARY", "EXPERIENCE", "SKILLS", "EDUCATION"}
CV_SECTION_HEADERS = {
    "CONTACT INFORMATION",
    "PROFESSIONAL SUMMARY / PROFILE",
    "CORE COMPETENCIES / TECHNICAL SKILLS",
    "PROFESSIONAL EXPERIENCE",
    "PROJECTS",
    "PUBLICATIONS / RESEARCH",
    "EDUCATION",
    "CERTIFICATIONS & TRAINING",
    "AWARDS & HONORS",
    "PROFESSIONAL AFFILIATIONS",
}
BULLET_PREFIXES = ("- ", "• ", "* ", "– ", "— ")


def extract_company(text: str) -> str:
    for line in text.splitlines():
        if line.startswith("COMPANY_NAME:"):
            company = line.split(":", 1)[1].strip()
            # Sanitize for filename
            company = re.sub(r"[^\w\s-]", "", company).strip()
            company = re.sub(r"\s+", "_", company)
            return company or "Company"
    return "Company"


def extract_marked_body(text: str, start_marker: str, end_marker: str, label: str) -> str:
    if start_marker in text and end_marker in text:
        start = text.index(start_marker) + len(start_marker)
        end = text.index(end_marker)
        return text[start:end].strip()
    sys.exit(f"Required {label} markers not found in model output: {start_marker} ... {end_marker}")


def extract_resume_body(text: str) -> str:
    return extract_marked_body(text, "---RESUME_START---", "---RESUME_END---", "resume")


def extract_cv_body(text: str) -> str:
    return extract_marked_body(text, "---CV_START---", "---CV_END---", "cv")


def is_bullet(line: str) -> bool:
    return any(line.startswith(p) for p in BULLET_PREFIXES)


def strip_bullet(line: str) -> str:
    for p in BULLET_PREFIXES:
        if line.startswith(p):
            return line[len(p):].strip()
    return line.strip()


def parse_resume(body: str) -> dict:
    """Parse the LLM output into structured sections."""
    lines = [ln.rstrip() for ln in body.splitlines()]
    # Drop leading/trailing blanks
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()

    if len(lines) < 2:
        sys.exit("Parsed resume body is too short — model output may be malformed.")

    parsed = {
        "name": lines[0].strip().lstrip("*").rstrip("*").strip(),
        "contact": lines[1].strip(),
        "sections": {},  # header -> list[str]
    }

    current_header = None
    current_lines: list[str] = []

    for line in lines[2:]:
        stripped = line.strip()
        if stripped.upper() in SECTION_HEADERS:
            if current_header is not None:
                parsed["sections"][current_header] = current_lines
            current_header = stripped.upper()
            current_lines = []
        elif current_header is not None:
            current_lines.append(line)

    if current_header is not None:
        parsed["sections"][current_header] = current_lines

    return parsed


def validate_required_sections(parsed: dict, required: set[str], label: str) -> None:
    missing = [section for section in required if section not in parsed["sections"]]
    if missing:
        sys.exit(f"{label} is missing required sections: {', '.join(sorted(missing))}")


# ─── DOCX WRITING ────────────────────────────────────────────────────────────

def add_hyperlink(paragraph, url: str, text: str, font_size: int = 10):
    """Add a clickable hyperlink to a paragraph (python-docx has no native helper)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))  # half-points
    rPr.append(sz)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def write_contact_line(doc, contact: str, linkedin_url: str | None):
    """Render the contact line, turning the literal word 'LinkedIn' into a hyperlink."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    parts = [seg.strip() for seg in contact.split("|")]
    for i, segment in enumerate(parts):
        if i > 0:
            sep = p.add_run("  |  ")
            sep.font.size = Pt(10)
        if segment.lower() == "linkedin" and linkedin_url:
            add_hyperlink(p, linkedin_url, "LinkedIn", font_size=10)
        else:
            run = p.add_run(segment)
            run.font.size = Pt(10)


def write_section_header(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # Bottom border for visual separation
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


def write_bullet(doc, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    if p.runs:
        p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(18)


def write_job_header(doc, line: str):
    """Render a job header like 'Title | Company, Location | Dates'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)

    parts = [seg.strip() for seg in line.split("|")]
    if len(parts) == 3:
        title, company, dates = parts
        run = p.add_run(f"{title} | {company}")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"   {dates}")
        run.italic = True
        run.font.size = Pt(10)
    else:
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(10)


def write_skills_line(doc, line: str):
    """Render 'Category: skill1, skill2' with the category bolded."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if ":" in line:
        category, skills = line.split(":", 1)
        bold_run = p.add_run(f"{category.strip()}: ")
        bold_run.bold = True
        bold_run.font.size = Pt(10)
        plain_run = p.add_run(skills.strip())
        plain_run.font.size = Pt(10)
    else:
        run = p.add_run(line)
        run.font.size = Pt(10)


def write_education_line(doc, line: str):
    """Education degree line — bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(line)
    run.bold = True
    run.font.size = Pt(10)


def write_coursework_line(doc, line: str):
    """Render 'Relevant Coursework: course1, course2' with the label bolded."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    if ":" in line:
        label, courses = line.split(":", 1)
        bold_run = p.add_run(f"{label.strip()}: ")
        bold_run.bold = True
        bold_run.italic = True
        bold_run.font.size = Pt(10)
        plain_run = p.add_run(courses.strip())
        plain_run.font.size = Pt(10)
    else:
        run = p.add_run(line)
        run.font.size = Pt(10)


def write_subrole_header(doc, line: str):
    """Render an academic sub-role like 'Teaching Assistant — ML & Data Mining | Jun 2022 – Aug 2023'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    parts = [seg.strip() for seg in line.split("|")]
    if len(parts) == 2:
        title, dates = parts
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(f"   {dates}")
        run.italic = True
        run.font.size = Pt(10)
    else:
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(10)


def write_plain(doc, line: str):
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def write_docx(parsed: dict, output_path: str, linkedin_url: str | None = None):
    doc = Document()

    # Tight margins for one-page resume
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Name (centered, large, bold)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(parsed["name"].upper())
    run.bold = True
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

    # Contact line (with LinkedIn hyperlink)
    write_contact_line(doc, parsed["contact"], linkedin_url)

    # Render sections in fixed order
    section_order = ["SUMMARY", "EXPERIENCE", "SKILLS", "EDUCATION"]
    sections = parsed["sections"]

    for header in section_order:
        if header not in sections:
            continue
        write_section_header(doc, header)
        body_lines = sections[header]

        for raw in body_lines:
            line = raw.strip()
            if not line:
                continue

            if is_bullet(line):
                write_bullet(doc, strip_bullet(line))
            elif header == "SUMMARY":
                write_plain(doc, line)
            elif header == "EXPERIENCE":
                # Job headers contain " | "
                if " | " in line:
                    write_job_header(doc, line)
                else:
                    write_plain(doc, line)
            elif header == "SKILLS":
                write_skills_line(doc, line)
            elif header == "EDUCATION":
                # Coursework: "Relevant Coursework: ..." (and tolerant variants)
                lower = line.lower()
                if lower.startswith("relevant coursework:") or lower.startswith("coursework:"):
                    write_coursework_line(doc, line)
                # Degree line: 3-part "Degree | School, Location | Year"
                elif line.count("|") >= 2:
                    write_education_line(doc, line)
                # Sub-role: 2-part "Title | Dates" (TA, research role, etc.)
                elif " | " in line:
                    write_subrole_header(doc, line)
                else:
                    write_plain(doc, line)
            else:
                write_plain(doc, line)

    doc.save(output_path)


def write_cv_docx(parsed: dict, output_path: str, linkedin_url: str | None = None):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(parsed["name"].upper())
    run.bold = True
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

    write_contact_line(doc, parsed["contact"], linkedin_url)

    section_order = [
        "CONTACT INFORMATION",
        "PROFESSIONAL SUMMARY / PROFILE",
        "CORE COMPETENCIES / TECHNICAL SKILLS",
        "PROFESSIONAL EXPERIENCE",
        "PROJECTS",
        "PUBLICATIONS / RESEARCH",
        "EDUCATION",
        "CERTIFICATIONS & TRAINING",
        "AWARDS & HONORS",
        "PROFESSIONAL AFFILIATIONS",
    ]
    sections = parsed["sections"]

    for header in section_order:
        if header not in sections:
            continue
        write_section_header(doc, header)
        body_lines = sections[header]

        for raw in body_lines:
            line = raw.strip()
            if not line:
                continue

            if is_bullet(line):
                write_bullet(doc, strip_bullet(line))
            elif "EXPERIENCE" in header and " | " in line:
                write_job_header(doc, line)
            elif "SKILLS" in header and ":" in line:
                write_skills_line(doc, line)
            elif header == "EDUCATION" and line.count("|") >= 2:
                write_education_line(doc, line)
            else:
                write_plain(doc, line)

    doc.save(output_path)



# ─── LOCAL DIFF SUMMARY (no second API call) ─────────────────────────────────

def _similarity(a: str, b: str) -> float:
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()


def _extract_bullets(sections: dict) -> dict:
    """Return {job_title: [bullet_text, ...]} for the EXPERIENCE section."""
    result = {}
    current_job = "Unknown Role"
    for line in sections.get("EXPERIENCE", []):
        line = line.strip()
        if not line:
            continue
        if " | " in line and not is_bullet(line):
            current_job = line.split("|")[0].strip()
        elif is_bullet(line):
            result.setdefault(current_job, []).append(strip_bullet(line))
    return result


def _extract_skills(sections: dict) -> set:
    skills = set()
    for line in sections.get("SKILLS", []):
        if ":" in line:
            _, _, vals = line.partition(":")
            for s in vals.split(","):
                s = s.strip()
                if s:
                    skills.add(s.lower())
    return skills


def _parse_original(resume_text: str) -> dict:
    """Best-effort parse of the plain-text original resume into sections."""
    sections = {}
    current = None
    HEADERS = {"SUMMARY", "PROFESSIONAL EXPERIENCE", "EXPERIENCE",
               "TECHNICAL SKILLS", "SKILLS", "EDUCATION"}
    for line in resume_text.splitlines():
        stripped = line.strip().upper().rstrip(":")
        if stripped in HEADERS:
            key = "EXPERIENCE" if "EXPERIENCE" in stripped else \
                  "SKILLS"     if "SKILL"      in stripped else stripped
            current = key
            sections.setdefault(current, [])
        elif current:
            sections[current].append(line)
    return {"sections": sections}


def compare_resumes(original_text: str, tailored_parsed: dict, jd_text: str, company: str) -> str:
    """Diff original vs tailored resume locally — zero API calls."""
    original_parsed = _parse_original(original_text)

    orig_bullets = _extract_bullets(original_parsed["sections"])
    tail_bullets = _extract_bullets(tailored_parsed["sections"])
    orig_skills  = _extract_skills(original_parsed["sections"])
    tail_skills  = _extract_skills(tailored_parsed["sections"])

    all_orig = [(job, b) for job, bullets in orig_bullets.items() for b in bullets]
    all_tail = [(job, b) for job, bullets in tail_bullets.items() for b in bullets]

    removed, modified, kept = [], [], []
    used = set()

    for job, orig_b in all_orig:
        best_score, best_idx, best_tail = 0.0, -1, ""
        for i, (_, tail_b) in enumerate(all_tail):
            if i in used:
                continue
            score = _similarity(orig_b, tail_b)
            if score > best_score:
                best_score, best_idx, best_tail = score, i, tail_b

        if best_score >= 0.85:
            kept.append(orig_b)
            used.add(best_idx)
        elif best_score >= 0.40:
            modified.append((job, orig_b, best_tail))
            used.add(best_idx)
        else:
            removed.append((job, orig_b))

    skills_removed = orig_skills - tail_skills
    skills_added   = tail_skills - orig_skills

    missing_terms = extract_missing_jd_terms(original_text, jd_text)
    tail_experience_text = "\n".join(tailored_parsed["sections"].get("EXPERIENCE", [])).lower()
    contextually_added_terms = [
        term for term in missing_terms if _contains_term(tail_experience_text, term)
    ]

    orig_edu = [l.strip() for l in original_parsed["sections"].get("EDUCATION", []) if l.strip()]
    tail_edu = [l.strip() for l in tailored_parsed["sections"].get("EDUCATION", []) if l.strip()]
    edu_added = [l for l in tail_edu if not any(_similarity(l, o) > 0.7 for o in orig_edu)]

    jd_words = set(re.findall(r'\b[A-Za-z][\w.+/-]{2,}\b', jd_text))
    tail_flat = " ".join(line for lines in tailored_parsed["sections"].values() for line in lines).lower()
    matched_kw = sorted(w for w in jd_words if w.lower() in tail_flat)
    missed_kw  = sorted(w for w in jd_words
                        if w.lower() not in tail_flat and len(w) > 4 and w[0].isupper())

    out = [
        f"Resume Tailoring Summary — {company}",
        f"Generated: {datetime.now():%Y-%m-%d %H:%M}",
        "",
        "─" * 60,
        f"BULLETS REMOVED  ({len(removed)})",
        "─" * 60,
    ]
    if removed:
        for job, b in removed:
            out += [f"  [{job}]", f"    - {b}"]
    else:
        out.append("  None")

    out += ["", "─" * 60, f"BULLETS MODIFIED  ({len(modified)})", "─" * 60]
    if modified:
        for job, orig_b, tail_b in modified:
            out += [f"  [{job}]", f"    Before: {orig_b}", f"    After:  {tail_b}"]
    else:
        out.append("  None")

    out += ["", "─" * 60, f"BULLETS KEPT UNCHANGED  ({len(kept)})", "─" * 60]
    for b in kept:
        out.append(f"  - {b}")

    out += ["", "─" * 60, "SKILLS REMOVED", "─" * 60]
    out.append("  " + (", ".join(sorted(skills_removed)) if skills_removed else "None"))

    out += ["", "─" * 60, "SKILLS ADDED", "─" * 60]
    out.append("  " + (", ".join(sorted(skills_added)) if skills_added else "None"))

    out += ["", "─" * 60, "JD TERMS ADDED IN EXPERIENCE (context check)", "─" * 60]
    out.append("  " + (", ".join(contextually_added_terms) if contextually_added_terms else "None"))

    out += ["", "─" * 60, "EDUCATION — CONTENT ADDED", "─" * 60]
    if edu_added:
        for l in edu_added:
            out.append(f"  + {l}")
    else:
        out.append("  No new content added")

    out += ["", "─" * 60, "JD KEYWORDS MATCHED", "─" * 60]
    out.append("  " + (", ".join(matched_kw[:20]) if matched_kw else "None"))

    out += ["", "─" * 60, "JD KEYWORDS POSSIBLY MISSED", "─" * 60]
    out.append("  " + (", ".join(missed_kw[:15]) if missed_kw else "None"))
    out.append("  (Review — may represent genuine skill gaps for this JD)")

    return "\n".join(out)


def write_summary_docx(summary_text: str, output_path: str, company: str):
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin   = Pt(72)
        section.right_margin  = Pt(72)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    for line in summary_text.splitlines():
        s = line.rstrip()
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)

        if s.startswith("Resume Tailoring Summary"):
            r = p.add_run(s); r.bold = True; r.font.size = Pt(13)
        elif s.startswith("─"):
            r = p.add_run(s)
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            r.font.size = Pt(9)
        elif re.match(r'^[A-Z ]+\([0-9]+\)$', s.strip()) or (s.isupper() and s.strip("─ ")):
            r = p.add_run(s); r.bold = True; r.font.size = Pt(10)
        elif s.strip().startswith("Before:"):
            r = p.add_run(s)
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            r.font.size = Pt(9)
        elif s.strip().startswith("After:"):
            r = p.add_run(s)
            r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
            r.font.size = Pt(9)
        elif "(Review" in s:
            r = p.add_run(s); r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x80, 0x40, 0x00)
        else:
            r = p.add_run(s); r.font.size = Pt(10)

    doc.save(output_path)


# ─── MAIN ────────────────────────────────────────────────────────────────────

def main():
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load optional local secrets file (kept out of git).
    load_dotenv_file(DOTENV_PATH)

    resume_text, jd_text = load_inputs()
    print(f"Resume: {len(resume_text)} chars | JD: {len(jd_text)} chars")

    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip().strip('"').strip("'")
    if not api_key:
        sys.exit(
            "ANTHROPIC_API_KEY is not set.\n"
            "Add it to your shell env or create a .env file in this project with:\n"
            "ANTHROPIC_API_KEY=sk-ant-..."
        )

    system_prompt = load_system_prompt(RULES_PATH)

    try:
        missing_jd_terms = extract_missing_jd_terms(resume_text, jd_text)
        if missing_jd_terms:
            print("Possible JD skill/tool gaps detected:", ", ".join(missing_jd_terms))
        output, usage = call_claude(
            resume_text,
            jd_text,
            api_key,
            system_prompt,
            missing_jd_terms=missing_jd_terms,
        )
    except anthropic.APIError as e:
        sys.exit(f"API call failed: {e}")

    print(f"Tokens — in: {usage['input_tokens']}, out: {usage['output_tokens']}")

    company = extract_company(output)
    print(f"Company: {company}")

    body = extract_resume_body(output)
    parsed = parse_resume(body)

    validate_required_sections(parsed, SECTION_HEADERS, "Resume")

    if not parsed["sections"]:
        debug_path = os.path.join(OUTPUT_DIR, f"raw_output_{datetime.now():%Y%m%d_%H%M%S}.txt")
        with open(debug_path, "w", encoding="utf-8") as f:
            f.write(output)
        sys.exit(f"No sections parsed. Raw output saved to: {debug_path}")

    def safe_write(fn, preferred, *args, **kwargs):
        try:
            fn(*args, output_path=preferred, **kwargs)
            return preferred
        except PermissionError:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            base, ext = os.path.splitext(preferred)
            fallback = f"{base}_{ts}{ext}"
            print(f"\n⚠  {os.path.basename(preferred)} is open in Word — saving to {os.path.basename(fallback)}")
            fn(*args, output_path=fallback, **kwargs)
            return fallback

    resume_path = safe_write(
        write_docx,
        os.path.join(OUTPUT_DIR, f"tailored_resume_{company}.docx"),
        parsed,
        linkedin_url=LINKEDIN_URL,
    )
    print(f"\n✓ Resume:  {resume_path} ({os.path.getsize(resume_path)/1024:.1f} KB)")


    # Local diff — no second API call
    print("Comparing resumes...")
    summary_text = compare_resumes(resume_text, parsed, jd_text, company)

    # Print to console
    print("\n" + summary_text)

    # Also write to docx
    summary_path = safe_write(
        write_summary_docx,
        os.path.join(OUTPUT_DIR, f"changes_summary_{company}.docx"),
        summary_text,
        company=company,
    )
    print(f"\n✓ Summary: {summary_path} ({os.path.getsize(summary_path)/1024:.1f} KB)")


if __name__ == "__main__":
    main()

